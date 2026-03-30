from django.shortcuts import render
from .forms import UserRegistrationForm
from django.contrib import messages
from .models import UserRegistrationModel
from django.conf import settings
from django.conf import settings
import pandas as pd


#user registration
def UserRegisterActions(request):
    if request.method == 'POST':
        form = UserRegistrationForm(request.POST)
        if form.is_valid():
            print('Data is Valid')
            form.save()
            messages.success(request, 'You have been successfully registered')
            form = UserRegistrationForm()
            return render(request, 'UserRegistrations.html', {'form': form})
        else:
            messages.success(request, 'Email or Mobile Already Existed')
            print("Invalid form")
    else:
        form = UserRegistrationForm()
    return render(request, 'UserRegistrations.html', {'form': form})


#user login check
def UserLoginCheck(request):
    if request.method == "POST":
        loginid=request.POST.get("loginid")
        password=request.POST.get("pswd")
        print(loginid)
        print(password)
        try:
            check=UserRegistrationModel.objects.get(loginid=loginid,password=password)
            status=check.status
            if status=="activated":
                request.session['id']=check.id
                request.session['loginid']=check.loginid
                request.session['password']=check.password
                request.session['email']=check.email
                return render(request,'users/UserHome.html',{})
            else:
                messages.success(request,"your account not activated")
            return render(request,"UserLogin.html")
        except Exception as e:
            print('=======>',e)
        messages.success(request,'invalid details')
    return render(request,'UserLogin.html',{})


    
def UserHome(request):
    return render(request,"users/UserHome.html",{})


def application(request):
    applications = RankedResume.objects.filter(status='present')
    return render(request, 'users/applications.html', {"results": applications})

def past_applications(request):
    applications = RankedResume.objects.filter(status='past')
    return render(request, 'users/past_applications.html', {"results": applications})

def training(request):
    # Mock data for graph
    data = {
        'labels': ['Resumes Uploaded', 'High Score (>75%)', 'Offer Letters Sent'],
        'values': [
            RankedResume.objects.count(), 
            RankedResume.objects.filter(similarity_score__gt=75).count(),
            RankedResume.objects.filter(status='past').count()
        ]
    }
    return render(request, 'users/training.html', {'graph_data': data})


from django.core.mail import EmailMessage
from django.conf import settings
from django.contrib import messages
from django.shortcuts import render
from docx import Document
from io import BytesIO
from django.utils import timezone
def create_offer_letter(candidate_name, candidate_email, candidate_phone, offer_letter_path):
    """
    Function to dynamically create a customized offer letter using the candidate's details.
    """
    # Create a new Word document (DOCX format)
    doc = Document()
    doc.add_paragraph("XYZ COMPANY")
    doc.add_paragraph("123 Business Road, Suite 100")
    doc.add_paragraph("City, State, Zip Code")
    doc.add_paragraph("Phone: (123) 456-7890")
    doc.add_paragraph("Email: hr@xyzcompany.com")
    doc.add_paragraph("Website: www.xyzcompany.com")
    doc.add_paragraph()

    doc.add_paragraph(f"Date: {timezone.now().strftime('%B %d, %Y')}")
    doc.add_paragraph()

    # Candidate Information
    doc.add_paragraph(f"Dear {candidate_name},")
    doc.add_paragraph(f"Email: {candidate_email}")
    doc.add_paragraph(f"Phone: {candidate_phone}")
    doc.add_paragraph()

    # Offer Letter Content
    doc.add_paragraph(
        f"We are pleased to offer you the position at XYZ Company. After careful consideration, we believe your skills "
        f"and experience align well with our organization's goals, and we are excited to have you join our team."
    )

    doc.add_paragraph("Please review the offer carefully and confirm your acceptance by signing and returning the attached acceptance form.")
    doc.add_paragraph("We look forward to your response and are excited about the possibility of working together.")
    doc.add_paragraph()
    doc.add_paragraph("If you have any questions, please feel free to reach out to me directly at hr@xyzcompany.com.")
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("[Your Name]")
    doc.add_paragraph("[Your Job Title]")
    doc.add_paragraph("XYZ Company")

    doc.add_paragraph("---")
    doc.add_paragraph("Acceptance of Offer:")
    doc.add_paragraph(f"I, {candidate_name}, hereby accept the offer extended by XYZ Company for the position as outlined in the above letter.")
    doc.add_paragraph("I agree to the terms and conditions of employment.")
    doc.add_paragraph("Signature: _________________________")
    doc.add_paragraph("Date: _________________________")
    doc.add_paragraph("Please return this signed acceptance to HR at hr@xyzcompany.com.")

    # Save the offer letter to a temporary file
    offer_letter_dir = os.path.join(settings.MEDIA_ROOT, "offer_letters")

# Create the directory if it doesn't exist
    if not os.path.exists(offer_letter_dir):
        os.makedirs(offer_letter_dir)

    # Define the file path
    offer_letter_path = os.path.join(offer_letter_dir, f"offer_letter_{candidate_email}.docx")

    # Save the offer letter
    doc.save(offer_letter_path)

    return offer_letter_path


def send_offer_letter_to_candidate(request):
    """
    Django view to handle the POST request and send a custom offer letter to the selected candidate.
    """
    if request.method == 'POST':
        # Get candidate's details from the POST request
        candidate_name = request.POST.get('name')
        candidate_email = request.POST.get('email')
        candidate_phone = '8520852085'

        if not candidate_name or not candidate_email or not candidate_phone:
            messages.error(request, "All fields are required.")
            return render(request, 'users/applications.html')

        # Generate the offer letter
        offer_letter_path = create_offer_letter(candidate_name, candidate_email, candidate_phone, None)

        # Send the offer letter
        subject = 'Job Offer: Congratulations on Your Selection'
        message = (
            f"Dear {candidate_name},\n\n"
            "We are pleased to inform you that you have been selected for the position at XYZ Company.\n"
            "Please find attached your official offer letter, including details about your role, salary, and benefits.\n\n"
            "We are excited to have you join our team and look forward to working with you.\n\n"
            "Best regards,\n"
            "XYZ Company\n"
            "Human Resources Team"
        )

        try:
            # Create the email with the attachment (offer letter)
            email_message = EmailMessage(
                subject=subject,
                body=message,
                from_email=settings.DEFAULT_FROM_EMAIL,
                to=[candidate_email]
            )
            email_message.attach_file(offer_letter_path)
            email_message.send(fail_silently=False)

            # Success message
            rr=RankedResume.objects.get(email=candidate_email)
            rr.status = 'past'
            rr.save()
            messages.success(request, f"Offer letter sent to {candidate_email} successfully!")
            applications=RankedResume.objects.filter(status='present')
            return render(request,'users/applications.html',{"results":applications})
        except Exception as e:
            # Error message if something goes wrong
            messages.error(request, f"Failed to send offer letter to {candidate_email}: {str(e)}")

        applications=RankedResume.objects.all()
        return render(request,'users/applications.html',{"results":applications})


    # If the request is GET or invalid, return the form again
    applications=RankedResume.objects.all()
    return render(request,'users/applications.html',{"results":applications})


# views.py
import os
import re
from django.shortcuts import render
from django.core.files.storage import FileSystemStorage
from django.conf import settings
from .models import RankedResume
import PyPDF2
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Extract text from PDFs
def extract_text_from_pdf(pdf_path):
    with open(pdf_path, "rb") as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""  # Handle None case
        return text
import os
import re
import PyPDF2
from django.shortcuts import render
from django.core.files.storage import FileSystemStorage
from django.conf import settings
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from .models import RankedResume


# =====================================
# Extract Text from PDF
# =====================================
def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with open(pdf_path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
    except:
        pass
    return text


# =====================================
# Extract Email & Phone
# =====================================
def extract_entities(text):
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'
    phone_pattern = r'\b\d{10}\b'

    emails = re.findall(email_pattern, text)
    phones = re.findall(phone_pattern, text)

    email = emails[0] if emails else "Not Found"
    phone = phones[0] if phones else "Not Found"

    return email, phone


# =====================================
# Resume Ranking + Save to DB
# =====================================
def index(request):

    if request.method == 'POST':

        job_description = request.POST.get('job_description')
        resume_files = request.FILES.getlist('resume_files')

        upload_dir = os.path.join(settings.MEDIA_ROOT, "uploads")
        os.makedirs(upload_dir, exist_ok=True)

        fs = FileSystemStorage(location=upload_dir)

        processed_resumes = []

        # -----------------------------
        # Save & Process Resumes
        # -----------------------------
        for resume_file in resume_files:

            if not resume_file.name.lower().endswith('.pdf'):
                continue

            filename = fs.save(resume_file.name, resume_file)
            resume_path = os.path.join(upload_dir, filename)

            resume_text = extract_text_from_pdf(resume_path)
            email, phone = extract_entities(resume_text)

            name = os.path.splitext(resume_file.name)[0]

            processed_resumes.append(
                (name, email, phone, resume_text, resume_path)
            )

        # -----------------------------
        # TF-IDF + Cosine Similarity
        # -----------------------------
        tfidf_vectorizer = TfidfVectorizer()
        job_vector = tfidf_vectorizer.fit_transform([job_description])

        ranked_resumes = []

        for name, email, phone, resume_text, resume_path in processed_resumes:

            resume_vector = tfidf_vectorizer.transform([resume_text])

            similarity = cosine_similarity(
                job_vector,
                resume_vector
            )[0][0] * 100

            similarity = round(similarity, 2)

            ranked_resumes.append(
                (name, email, phone, similarity, resume_path)
            )

            # -----------------------------
            # SAVE TO DATABASE (ACTIVE)
            # -----------------------------
            if similarity > 20 and email != "Not Found":

                RankedResume.objects.update_or_create(
                    email=email,
                    defaults={
                        'name': name,
                        'similarity_score': similarity,
                        'status': 'present'
                    }
                )

        # -----------------------------
        # Sort by Highest Similarity
        # -----------------------------
        ranked_resumes.sort(key=lambda x: x[3], reverse=True)

        # -----------------------------
        # Prepare Results for Template
        # -----------------------------
        results = []

        for name, email, phone, similarity, resume_path in ranked_resumes:

            rel_path = os.path.relpath(resume_path, settings.MEDIA_ROOT)
            resume_url = settings.MEDIA_URL + rel_path.replace('\\', '/')

            results.append(
                (name, email, phone, similarity, resume_url)
            )

        return render(
            request,
            'users/prediction_results.html',
            {'results': results}
        )

    return render(request, 'users/upload_resumes.html')